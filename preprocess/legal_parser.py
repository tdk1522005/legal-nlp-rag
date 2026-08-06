import json
import re
import unicodedata
from collections import Counter
from pathlib import Path
from typing import Any, Iterator

from docx import Document


class LegalDocxParser:
    """
    Đọc văn bản pháp luật DOCX và chuyển thành cây:

    DOCUMENT
        PART
            CHAPTER
                SECTION
                    SUBSECTION
                        ARTICLE
                            CLAUSE
                                POINT

    Parser có thêm hai cơ chế bảo vệ:

    1. Nội dung nằm trong dấu ngoặc kép tiếng Việt “...” được giữ dưới
       dạng TEXT. Các số thứ tự bên trong đoạn trích dẫn không được nhận
       nhầm thành Điều, Khoản hoặc Điểm của văn bản đang phân tích.

    2. Phần chú thích cuối văn bản, thường bắt đầu bằng [1], [2], ...
       được tách riêng khỏi cây pháp luật chính.
    """

    PART_PATTERN = re.compile(
        r"^\s*PHẦN(?:\s+THỨ)?\s+(.+?)\s*$",
        flags=re.IGNORECASE,
    )

    CHAPTER_PATTERN = re.compile(
        r"^\s*CHƯƠNG\s+([IVXLCDM\d]+)\s*$",
        flags=re.IGNORECASE,
    )

    SUBSECTION_PATTERN = re.compile(
        r"^\s*TIỂU\s+MỤC\s+(\d+[A-Z]?)\s*$",
        flags=re.IGNORECASE,
    )

    SECTION_PATTERN = re.compile(
        r"^\s*MỤC\s+(\d+[A-Z]?)\s*$",
        flags=re.IGNORECASE,
    )

    ARTICLE_PATTERN = re.compile(
        r"^\s*Điều\s+(\d+[a-zA-Z]?)"
        r"\s*[.:\-]?\s*(.*)$",
        flags=re.IGNORECASE,
    )

    CLAUSE_PATTERN = re.compile(
        r"^\s*(\d+)\.\s*(.+)$",
    )

    POINT_PATTERN = re.compile(
        r"^\s*([a-zA-ZđĐ])\)\s*(.+)$",
    )

    # Ví dụ:
    #   Điều 58 của Luật Quy hoạch đô thị và nông thôn ...
    # Đây là câu dẫn chiếu đến văn bản khác, không phải tiêu đề Điều
    # của văn bản đang được parse.
    ARTICLE_REFERENCE_PATTERN = re.compile(
        r"^\s*Điều\s+\d+[a-zA-Z]?"
        r"\s+của\s+"
        r"(?:Bộ\s+luật|Luật|Nghị\s+quyết|Pháp\s+lệnh|"
        r"Nghị\s+định|Thông\s+tư|Văn\s+bản)\b",
        flags=re.IGNORECASE,
    )

    # Các văn bản hợp nhất thường đặt toàn bộ chú thích ở cuối file.
    # Chú thích bắt đầu bằng [1], [2], [25], ...
    ENDNOTE_PATTERN = re.compile(
        r"^\s*\[\d+\](?:\s|$)",
    )

    STRUCTURE_TYPES = {
        "PART",
        "CHAPTER",
        "SECTION",
        "SUBSECTION",
    }

    OPEN_QUOTE = "“"
    CLOSE_QUOTE = "”"

    def __init__(self) -> None:
        self.law_id = ""
        self.stats: Counter[str] = Counter()

    # =====================================================
    # 1. CHUẨN HÓA
    # =====================================================

    @staticmethod
    def normalize_text(text: str) -> str:
        """
        Xóa khoảng trắng dư thừa nhưng giữ nguyên nội dung.
        """
        return re.sub(
            r"\s+",
            " ",
            text,
        ).strip()

    @staticmethod
    def slugify(value: str) -> str:
        """
        Chuyển chuỗi thành dạng phù hợp cho node_id.
        """
        value = unicodedata.normalize(
            "NFD",
            value,
        )

        value = "".join(
            character
            for character in value
            if unicodedata.category(character) != "Mn"
        )

        value = value.lower()
        value = value.replace("đ", "d")

        value = re.sub(
            r"[^a-z0-9]+",
            "-",
            value,
        )

        return value.strip("-")

    # =====================================================
    # 2. NHẬN DIỆN ĐOẠN ĐẶC BIỆT
    # =====================================================

    @classmethod
    def is_article_reference(cls, text: str) -> bool:
        """
        Trả về True khi paragraph chỉ đang dẫn chiếu đến Điều của
        một văn bản khác.

        Ví dụ:
            Điều 58 của Luật Quy hoạch đô thị và nông thôn ...
        """
        return bool(
            cls.ARTICLE_REFERENCE_PATTERN.match(
                cls.normalize_text(text)
            )
        )

    @classmethod
    def is_endnote_start(cls, text: str) -> bool:
        """
        Nhận diện paragraph mở đầu phần chú thích cuối văn bản.
        """
        return bool(
            cls.ENDNOTE_PATTERN.match(
                cls.normalize_text(text)
            )
        )

    @classmethod
    def update_quote_depth(
        cls,
        text: str,
        current_depth: int,
    ) -> int:
        """
        Cập nhật trạng thái khối trích dẫn dùng dấu “ và ”.

        Dùng depth thay vì boolean để không vỡ khi văn bản có trích dẫn
        lồng nhau. Giá trị không bao giờ nhỏ hơn 0.
        """
        depth = current_depth

        for character in text:
            if character == cls.OPEN_QUOTE:
                depth += 1
            elif character == cls.CLOSE_QUOTE:
                depth = max(0, depth - 1)

        return depth

    # =====================================================
    # 3. PHÂN LOẠI PARAGRAPH
    # =====================================================

    def classify(self, text: str) -> str:
        text = self.normalize_text(text)

        if not text:
            return "EMPTY"

        if self.PART_PATTERN.match(text):
            return "PART"

        if self.CHAPTER_PATTERN.match(text):
            return "CHAPTER"

        if self.SUBSECTION_PATTERN.match(text):
            return "SUBSECTION"

        if self.SECTION_PATTERN.match(text):
            return "SECTION"

        # Phải kiểm tra dẫn chiếu trước ARTICLE_PATTERN.
        if self.is_article_reference(text):
            return "TEXT"

        if self.ARTICLE_PATTERN.match(text):
            return "ARTICLE"

        if self.CLAUSE_PATTERN.match(text):
            return "CLAUSE"

        if self.POINT_PATTERN.match(text):
            return "POINT"

        return "TEXT"

    # =====================================================
    # 4. TẠO NODE
    # =====================================================

    def create_node(
        self,
        node_type: str,
        number: str | None,
        heading: str,
        title: str,
        paragraph_index: int,
        parent: dict[str, Any],
    ) -> dict[str, Any]:
        parent_id = parent["node_id"]

        normalized_number = self.slugify(
            number or str(paragraph_index)
        )

        if node_type == "ARTICLE":
            node_id = (
                f"{self.law_id}/article/"
                f"{normalized_number}"
            )

        elif node_type == "CLAUSE":
            node_id = (
                f"{parent_id}/clause/"
                f"{normalized_number}"
            )

        elif node_type == "POINT":
            node_id = (
                f"{parent_id}/point/"
                f"{normalized_number}"
            )

        else:
            node_id = (
                f"{parent_id}/"
                f"{node_type.lower()}/"
                f"{normalized_number}"
            )

        return {
            "node_id": node_id,
            "node_type": node_type,
            "number": number,
            "heading": heading,
            "title": title,
            "paragraph_index": paragraph_index,
            "paragraphs": [],
            "children": [],
        }

    # =====================================================
    # 5. XÁC ĐỊNH NODE CHỨA NỘI DUNG
    # =====================================================

    @staticmethod
    def deepest_content_node(
        state: dict[str, dict[str, Any] | None],
        root: dict[str, Any],
    ) -> dict[str, Any]:
        """
        TEXT sẽ được gắn vào node gần nhất.

        Ưu tiên:
            POINT → CLAUSE → ARTICLE →
            SUBSECTION → SECTION → CHAPTER → PART
        """
        for key in (
            "point",
            "clause",
            "article",
            "subsection",
            "section",
            "chapter",
            "part",
        ):
            node = state.get(key)

            if node is not None:
                return node

        return root

    # =====================================================
    # 6. PARSE DOCX
    # =====================================================

    def parse(
        self,
        file_path: str | Path,
        law_id: str,
        document_title: str | None = None,
    ) -> dict[str, Any]:
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(
                f"Không tìm thấy file: {file_path}"
            )

        if file_path.suffix.lower() != ".docx":
            raise ValueError(
                "LegalDocxParser chỉ hỗ trợ file .docx."
            )

        self.law_id = law_id
        self.stats = Counter()

        document = Document(file_path)

        root: dict[str, Any] = {
            "node_id": law_id,
            "node_type": "DOCUMENT",
            "number": None,
            "heading": document_title or file_path.stem,
            "title": document_title or file_path.stem,
            "paragraph_index": None,
            "paragraphs": [],
            "children": [],
        }

        state: dict[str, dict[str, Any] | None] = {
            "part": None,
            "chapter": None,
            "section": None,
            "subsection": None,
            "article": None,
            "clause": None,
            "point": None,
        }

        pending_title_node: dict[str, Any] | None = None

        # depth > 0 nghĩa là paragraph hiện tại đang nằm trong khối
        # văn bản trích dẫn bằng dấu “...”.
        quote_depth = 0

        # Khi gặp [1], [2], ... sau phần thân luật, toàn bộ paragraph
        # còn lại được lưu vào endnotes và không tạo node pháp luật.
        in_endnotes = False
        endnotes: list[dict[str, Any]] = []

        for paragraph_index, paragraph in enumerate(
            document.paragraphs
        ):
            text = self.normalize_text(
                paragraph.text
            )

            if not text:
                self.stats["EMPTY"] += 1
                continue

            # -------------------------------------------------
            # PHẦN CHÚ THÍCH CUỐI VĂN BẢN
            # -------------------------------------------------

            # Chỉ kích hoạt sau khi đã đọc ít nhất một Điều để tránh
            # nhận nhầm nội dung ở phần mở đầu tài liệu.
            if (
                not in_endnotes
                and state["article"] is not None
                and self.is_endnote_start(text)
            ):
                in_endnotes = True
                pending_title_node = None
                quote_depth = 0

            if in_endnotes:
                endnotes.append(
                    {
                        "paragraph_index": paragraph_index,
                        "text": text,
                    }
                )
                self.stats["ENDNOTE"] += 1
                continue

            # -------------------------------------------------
            # KHỐI TRÍCH DẪN
            # -------------------------------------------------

            inside_quote_before = quote_depth > 0
            starts_with_open_quote = text.lstrip().startswith(
                self.OPEN_QUOTE
            )

            # Cập nhật trước các nhánh continue. Việc ép paragraph
            # thành TEXT vẫn dựa trên trạng thái trước khi cập nhật.
            quote_depth = self.update_quote_depth(
                text=text,
                current_depth=quote_depth,
            )

            if inside_quote_before or starts_with_open_quote:
                paragraph_type = "TEXT"
                self.stats["QUOTED_TEXT"] += 1
            else:
                paragraph_type = self.classify(text)

                if self.is_article_reference(text):
                    self.stats["ARTICLE_REFERENCE_TEXT"] += 1

            self.stats[paragraph_type] += 1

            # -------------------------------------------------
            # Tiêu đề của Phần, Chương, Mục hoặc Tiểu mục
            #
            # Ví dụ:
            #   Chương VIII
            #   GIAO DỊCH DÂN SỰ
            # -------------------------------------------------

            if (
                paragraph_type == "TEXT"
                and pending_title_node is not None
            ):
                pending_title_node["title"] = text
                pending_title_node = None
                continue

            if paragraph_type != "TEXT":
                pending_title_node = None

            # -------------------------------------------------
            # PHẦN
            # -------------------------------------------------

            if paragraph_type == "PART":
                match = self.PART_PATTERN.match(text)

                number = (
                    match.group(1).strip()
                    if match
                    else text
                )

                node = self.create_node(
                    node_type="PART",
                    number=number,
                    heading=text,
                    title="",
                    paragraph_index=paragraph_index,
                    parent=root,
                )

                root["children"].append(node)

                state["part"] = node
                state["chapter"] = None
                state["section"] = None
                state["subsection"] = None
                state["article"] = None
                state["clause"] = None
                state["point"] = None

                pending_title_node = node
                continue

            # -------------------------------------------------
            # CHƯƠNG
            # -------------------------------------------------

            if paragraph_type == "CHAPTER":
                match = self.CHAPTER_PATTERN.match(text)

                number = (
                    match.group(1).strip()
                    if match
                    else text
                )

                parent = state["part"] or root

                node = self.create_node(
                    node_type="CHAPTER",
                    number=number,
                    heading=text,
                    title="",
                    paragraph_index=paragraph_index,
                    parent=parent,
                )

                parent["children"].append(node)

                state["chapter"] = node
                state["section"] = None
                state["subsection"] = None
                state["article"] = None
                state["clause"] = None
                state["point"] = None

                pending_title_node = node
                continue

            # -------------------------------------------------
            # MỤC
            # -------------------------------------------------

            if paragraph_type == "SECTION":
                match = self.SECTION_PATTERN.match(text)

                number = (
                    match.group(1).strip()
                    if match
                    else text
                )

                parent = (
                    state["chapter"]
                    or state["part"]
                    or root
                )

                node = self.create_node(
                    node_type="SECTION",
                    number=number,
                    heading=text,
                    title="",
                    paragraph_index=paragraph_index,
                    parent=parent,
                )

                parent["children"].append(node)

                state["section"] = node
                state["subsection"] = None
                state["article"] = None
                state["clause"] = None
                state["point"] = None

                pending_title_node = node
                continue

            # -------------------------------------------------
            # TIỂU MỤC
            # -------------------------------------------------

            if paragraph_type == "SUBSECTION":
                match = self.SUBSECTION_PATTERN.match(text)

                number = (
                    match.group(1).strip()
                    if match
                    else text
                )

                parent = (
                    state["section"]
                    or state["chapter"]
                    or state["part"]
                    or root
                )

                node = self.create_node(
                    node_type="SUBSECTION",
                    number=number,
                    heading=text,
                    title="",
                    paragraph_index=paragraph_index,
                    parent=parent,
                )

                parent["children"].append(node)

                state["subsection"] = node
                state["article"] = None
                state["clause"] = None
                state["point"] = None

                pending_title_node = node
                continue

            # -------------------------------------------------
            # ĐIỀU
            # -------------------------------------------------

            if paragraph_type == "ARTICLE":
                match = self.ARTICLE_PATTERN.match(text)

                if match is None:
                    continue

                number = match.group(1).strip()
                title = match.group(2).strip()

                parent = (
                    state["subsection"]
                    or state["section"]
                    or state["chapter"]
                    or state["part"]
                    or root
                )

                node = self.create_node(
                    node_type="ARTICLE",
                    number=number,
                    heading=text,
                    title=title,
                    paragraph_index=paragraph_index,
                    parent=parent,
                )

                parent["children"].append(node)

                state["article"] = node
                state["clause"] = None
                state["point"] = None
                continue

            # -------------------------------------------------
            # KHOẢN
            # -------------------------------------------------

            if paragraph_type == "CLAUSE":
                match = self.CLAUSE_PATTERN.match(text)

                if match is None:
                    continue

                article = state["article"]

                if article is None:
                    root["paragraphs"].append(text)
                    self.stats["ORPHAN_CLAUSE"] += 1
                    continue

                number = match.group(1).strip()

                node = self.create_node(
                    node_type="CLAUSE",
                    number=number,
                    heading=text,
                    title="",
                    paragraph_index=paragraph_index,
                    parent=article,
                )

                node["paragraphs"].append(text)
                article["children"].append(node)

                state["clause"] = node
                state["point"] = None
                continue

            # -------------------------------------------------
            # ĐIỂM
            # -------------------------------------------------

            if paragraph_type == "POINT":
                match = self.POINT_PATTERN.match(text)

                if match is None:
                    continue

                parent = (
                    state["clause"]
                    or state["article"]
                )

                if parent is None:
                    root["paragraphs"].append(text)
                    self.stats["ORPHAN_POINT"] += 1
                    continue

                number = match.group(1).lower()

                node = self.create_node(
                    node_type="POINT",
                    number=number,
                    heading=text,
                    title="",
                    paragraph_index=paragraph_index,
                    parent=parent,
                )

                node["paragraphs"].append(text)
                parent["children"].append(node)

                state["point"] = node
                continue

            # -------------------------------------------------
            # TEXT KHÔNG ĐÁNH SỐ
            # -------------------------------------------------

            content_node = self.deepest_content_node(
                state=state,
                root=root,
            )

            content_node["paragraphs"].append(text)

        return {
            "schema_version": "1.1",
            "law_id": law_id,
            "document_title": (
                document_title or file_path.stem
            ),
            "source_file": str(file_path),
            "statistics": {
                "paragraphs_total": len(
                    document.paragraphs
                ),
                "tables_total": len(document.tables),
                "endnotes_total": len(endnotes),
                **dict(self.stats),
            },
            "endnotes": endnotes,
            "tree": root,
        }

    # =====================================================
    # 7. DUYỆT CÂY
    # =====================================================

    @staticmethod
    def iter_nodes(
        node: dict[str, Any],
    ) -> Iterator[dict[str, Any]]:
        yield node

        for child in node.get("children", []):
            yield from LegalDocxParser.iter_nodes(
                child
            )

    @staticmethod
    def find_nodes(
        tree: dict[str, Any],
        node_type: str,
    ) -> list[dict[str, Any]]:
        return [
            node
            for node in LegalDocxParser.iter_nodes(tree)
            if node.get("node_type") == node_type
        ]

    @staticmethod
    def find_article(
        tree: dict[str, Any],
        article_number: str,
    ) -> dict[str, Any] | None:
        for node in LegalDocxParser.iter_nodes(tree):
            if (
                node.get("node_type") == "ARTICLE"
                and str(node.get("number"))
                == str(article_number)
            ):
                return node

        return None

    # =====================================================
    # 8. XUẤT JSON
    # =====================================================

    @staticmethod
    def save_json(
        parsed_document: dict[str, Any],
        output_path: str | Path,
    ) -> None:
        output_path = Path(output_path)

        output_path.parent.mkdir(
            parents=True,
            exist_ok=True,
        )

        with output_path.open(
            mode="w",
            encoding="utf-8",
        ) as file:
            json.dump(
                parsed_document,
                file,
                ensure_ascii=False,
                indent=2,
            )