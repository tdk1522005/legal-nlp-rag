import argparse
import sys
from collections import defaultdict
from pathlib import Path
from typing import Any, Iterator

from docx import Document


RAG_MODEL_DIR = Path(__file__).resolve().parents[1]

if str(RAG_MODEL_DIR) not in sys.path:
    sys.path.insert(0, str(RAG_MODEL_DIR))


from preprocess.legal_parser import LegalDocxParser


DATA_DIR = RAG_MODEL_DIR / "data"

STRUCTURE_TYPES = {
    "PART",
    "CHAPTER",
    "SECTION",
    "SUBSECTION",
}


def normalize_path_label(
    node: dict[str, Any],
) -> str:
    heading = str(
        node.get("heading", "")
    ).strip()

    title = str(
        node.get("title", "")
    ).strip()

    if heading and title:
        if title.casefold() not in heading.casefold():
            return f"{heading} - {title}"

    return heading or title


def iter_articles_with_path(
    node: dict[str, Any],
    ancestors: list[dict[str, Any]] | None = None,
) -> Iterator[
    tuple[
        dict[str, Any],
        list[dict[str, Any]],
    ]
]:
    if ancestors is None:
        ancestors = []

    node_type = node.get("node_type")

    current_ancestors = list(ancestors)

    if node_type in STRUCTURE_TYPES:
        current_ancestors.append(node)

    if node_type == "ARTICLE":
        yield node, ancestors

    for child in node.get("children", []):
        yield from iter_articles_with_path(
            child,
            current_ancestors,
        )


def print_paragraph_context(
    document: Document,
    paragraph_index: int,
    context_size: int,
) -> None:
    start = max(
        paragraph_index - context_size,
        0,
    )

    end = min(
        paragraph_index + context_size + 1,
        len(document.paragraphs),
    )

    print(
        f"    Ngữ cảnh paragraph "
        f"{start} → {end - 1}:"
    )

    for index in range(start, end):
        text = " ".join(
            document.paragraphs[index]
            .text
            .split()
        )

        marker = (
            ">>>"
            if index == paragraph_index
            else "   "
        )

        if not text:
            text = "[EMPTY]"

        if len(text) > 220:
            text = text[:220] + "..."

        print(
            f"    {marker} [{index:04d}] {text}"
        )


def print_document_boundaries(
    document: Document,
    count: int = 15,
) -> None:
    non_empty = [
        (
            index,
            " ".join(paragraph.text.split()),
        )
        for index, paragraph in enumerate(
            document.paragraphs
        )
        if paragraph.text.strip()
    ]

    print("\n" + "=" * 90)
    print("PHẦN ĐẦU VĂN BẢN")
    print("=" * 90)

    for index, text in non_empty[:count]:
        print(f"[{index:04d}] {text[:220]}")

    print("\n" + "=" * 90)
    print("PHẦN CUỐI VĂN BẢN")
    print("=" * 90)

    for index, text in non_empty[-count:]:
        print(f"[{index:04d}] {text[:220]}")


def diagnose(
    file_path: Path,
    law_id: str,
    context_size: int,
) -> None:
    if not file_path.exists():
        raise FileNotFoundError(
            f"Không tìm thấy file: {file_path}"
        )

    document = Document(file_path)

    parser = LegalDocxParser()

    parsed_document = parser.parse(
        file_path=file_path,
        law_id=law_id,
        document_title=file_path.stem,
    )

    tree = parsed_document["tree"]

    article_records = list(
        iter_articles_with_path(tree)
    )

    articles_by_number: dict[
        str,
        list[
            tuple[
                dict[str, Any],
                list[dict[str, Any]],
            ]
        ],
    ] = defaultdict(list)

    for article, ancestors in article_records:
        article_number = str(
            article.get("number", "")
        )

        articles_by_number[
            article_number
        ].append(
            (
                article,
                ancestors,
            )
        )

    duplicates = {
        number: records
        for number, records
        in articles_by_number.items()
        if len(records) > 1
    }

    ordered_articles = sorted(
        article_records,
        key=lambda item: (
            item[0].get(
                "paragraph_index",
                0,
            )
        ),
    )

    print("=" * 90)
    print("CHẨN ĐOÁN VĂN BẢN PHÁP LUẬT")
    print("=" * 90)

    print(f"File: {file_path}")
    print(f"law_id: {law_id}")
    print(
        f"Tổng paragraph: "
        f"{len(document.paragraphs)}"
    )
    print(
        f"Tổng table: "
        f"{len(document.tables)}"
    )
    print(
        f"Số node Điều: "
        f"{len(article_records)}"
    )
    print(
        f"Số Điều duy nhất: "
        f"{len(articles_by_number)}"
    )
    print(
        f"Số số Điều bị lặp: "
        f"{len(duplicates)}"
    )

    print("\n" + "=" * 90)
    print("10 ĐIỀU ĐẦU TIÊN")
    print("=" * 90)

    for article, _ in ordered_articles[:10]:
        print(
            f"[{article['paragraph_index']:04d}] "
            f"{article['heading']}"
        )

    print("\n" + "=" * 90)
    print("15 ĐIỀU CUỐI CÙNG")
    print("=" * 90)

    for article, _ in ordered_articles[-15:]:
        print(
            f"[{article['paragraph_index']:04d}] "
            f"{article['heading']}"
        )

    if not duplicates:
        print("\n" + "=" * 90)
        print("KHÔNG PHÁT HIỆN SỐ ĐIỀU BỊ LẶP")
        print("=" * 90)

    else:
        print("\n" + "=" * 90)
        print("CÁC SỐ ĐIỀU BỊ LẶP")
        print("=" * 90)

        for article_number, records in sorted(
            duplicates.items(),
            key=lambda item: min(
                record[0]["paragraph_index"]
                for record in item[1]
            ),
        ):
            print(
                f"\nĐiều {article_number} "
                f"xuất hiện {len(records)} lần"
            )

            for occurrence, (
                article,
                ancestors,
            ) in enumerate(
                records,
                start=1,
            ):
                breadcrumb = [
                    normalize_path_label(node)
                    for node in ancestors
                    if normalize_path_label(node)
                ]

                print("\n" + "-" * 90)
                print(
                    f"Lần {occurrence}: "
                    f"paragraph "
                    f"{article['paragraph_index']}"
                )
                print(
                    f"Heading: "
                    f"{article['heading']}"
                )

                if breadcrumb:
                    print(
                        "Vị trí: "
                        + " > ".join(breadcrumb)
                    )
                else:
                    print(
                        "Vị trí: không có "
                        "Phần/Chương/Mục cha"
                    )

                print_paragraph_context(
                    document=document,
                    paragraph_index=article[
                        "paragraph_index"
                    ],
                    context_size=context_size,
                )

    print_document_boundaries(
        document=document,
        count=15,
    )


def resolve_file_path(
    file_argument: str,
) -> Path:
    path = Path(file_argument)

    if path.is_absolute():
        return path

    return DATA_DIR / path


def main() -> None:
    argument_parser = argparse.ArgumentParser()

    argument_parser.add_argument(
        "file",
        help=(
            "Đường dẫn DOCX tính từ "
            "thư mục data."
        ),
    )

    argument_parser.add_argument(
        "law_id",
        help="law_id dùng khi parse.",
    )

    argument_parser.add_argument(
        "--context",
        type=int,
        default=4,
        help=(
            "Số paragraph hiển thị trước "
            "và sau Điều bị lặp."
        ),
    )

    arguments = argument_parser.parse_args()

    try:
        diagnose(
            file_path=resolve_file_path(
                arguments.file
            ),
            law_id=arguments.law_id,
            context_size=arguments.context,
        )

    except Exception as error:
        print(f"[ERROR] {error}")
        sys.exit(1)


if __name__ == "__main__":
    main()